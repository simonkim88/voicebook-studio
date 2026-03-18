# document_parser.py - 문서 파싱 모듈
import os
import re
import json

# 문서 파서 임포트
try:
    import fitz
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False


class DocumentParser:
    """문서 파싱 클래스"""
    
    @staticmethod
    def parse_txt(filepath):
        """텍스트 파일 파싱"""
        encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        raise Exception("파일 인코딩을 감지할 수 없습니다.")
    
    @staticmethod
    def parse_rtf(filepath):
        """리치텍스트 파일 파싱 - 테이블 제외"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # RTF 테이블 제외 (\trowd ~ \row)
            content = re.sub(r'\\trowd.*?\\row', '', content, flags=re.DOTALL)
            
            # 기존 정리
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            content = re.sub(r'[{}]', '', content)
            content = re.sub(r'\\\*?\\[a-z]+\s?', '', content)
            content = re.sub(r'\\\d+\s?', '', content)
            return content.strip()
        except Exception as e:
            raise Exception(f"RTF 파싱 오류: {str(e)}")
    
    @staticmethod
    def parse_pdf(filepath):
        """PDF 파일 파싱 - 테이블 제외 시도"""
        if not PDF_AVAILABLE:
            raise Exception("PyMuPDF가 설치되어 있지 않습니다. 'pip install pymupdf'를 실행하세요.")
        try:
            doc = fitz.open(filepath)
            text = ""
            
            for page in doc:
                # 테이블 감지 및 제외 (PyMuPDF 1.23.5+)
                try:
                    tables = page.find_tables()
                    if tables and len(tables.tables) > 0:
                        # 테이블 영역을 제외한 텍스트 추출
                        table_bboxes = [table.bbox for table in tables.tables]
                        blocks = page.get_text("blocks")
                        for block in blocks:
                            block_bbox = fitz.Rect(block[:4])
                            # 테이블 영역과 겹치지 않는 블록만 포함
                            is_in_table = False
                            for table_bbox in table_bboxes:
                                if block_bbox.intersects(table_bbox):
                                    is_in_table = True
                                    break
                            if not is_in_table:
                                text += block[4] + "\n"
                    else:
                        text += page.get_text()
                except:
                    # 테이블 감지 실패 시 일반 텍스트 추출
                    text += page.get_text()
            
            doc.close()
            return text
        except Exception as e:
            raise Exception(f"PDF 파싱 오류: {str(e)}")
    
    @staticmethod
    def parse_docx(filepath):
        """Word 문서 파싱 - 테이블 제외"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx가 설치되어 있지 않습니다. 'pip install python-docx'를 실행하세요.")
        try:
            doc = Document(filepath)
            paragraphs = []
            
            for para in doc.paragraphs:
                # 테이블 안에 있는 단락은 제외
                if para._element.getparent().tag.endswith('tbl'):
                    continue
                if para.text.strip():
                    paragraphs.append(para.text)
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise Exception(f"DOCX 파싱 오류: {str(e)}")
    
    @staticmethod
    def parse_epub(filepath):
        """EPUB 전자책 파싱 - 테이블 제외"""
        if not EPUB_AVAILABLE:
            raise Exception("ebooklib이 설치되어 있지 않습니다. 'pip install ebooklib'를 실행하세요.")
        try:
            book = epub.read_epub(filepath)
            texts = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    try:
                        content = item.get_content().decode('utf-8')
                        
                        # HTML 테이블 제외 (<table> ~ </table>)
                        content = re.sub(r'<table[^>]*>.*?</table>', ' ', content, flags=re.DOTALL | re.IGNORECASE)
                        
                        # 나머지 HTML 태그 제거
                        content = re.sub(r'<[^>]+>', ' ', content)
                        content = re.sub(r'\s+', ' ', content).strip()
                        if content:
                            texts.append(content)
                    except:
                        continue
            return "\n\n".join(texts)
        except Exception as e:
            raise Exception(f"EPUB 파싱 오류: {str(e)}")
    
    @classmethod
    def parse(cls, filepath):
        """파일 확장자에 따라 자동 파싱"""
        ext = os.path.splitext(filepath)[1].lower()
        parsers = {
            '.txt': cls.parse_txt,
            '.rtf': cls.parse_rtf,
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.epub': cls.parse_epub
        }
        if ext in parsers:
            return parsers[ext](filepath)
        return cls.parse_txt(filepath)
    
    @classmethod
    def get_supported_extensions(cls):
        """지원하는 파일 확장자 목록"""
        exts = ['.txt']
        if PDF_AVAILABLE:
            exts.append('.pdf')
        if DOCX_AVAILABLE:
            exts.append('.docx')
        if EPUB_AVAILABLE:
            exts.append('.epub')
        exts.append('.rtf')
        return exts


# 상수
VOICE_OPTIONS = [
    ("Vivian", "Chinese", "밝고 약간 날카로운 젊은 여성 목소리"),
    ("Serena", "Chinese", "따뜻하고 부드러운 젊은 여성 목소리"),
    ("Uncle_Fu", "Chinese", "풍부한 저음의 중년 남성 목소리"),
    ("Dylan", "Chinese (Beijing Dialect)", "명랑하고 자연스러운 북경 사투리 남성 목소리"),
    ("Eric", "Chinese (Sichuan Dialect)", "쾌활하고 약간 쉰 듯한 청두 사투리 남성 목소리"),
    ("Ryan", "English", "강한 리듬감이 있는 역동적인 남성 목소리"),
    ("Aiden", "English", "맑은 중음역의 밝은 미국식 남성 목소리"),
    ("Ono_Anna", "Japanese", "경쾌하고 재치 있는 일본 여성 목소리"),
    ("Sohee", "Korean", "풍부한 감정의 따뜻한 한국 여성 목소리")
]

# 커스텀 음성 프리셋 (voices/ 디렉토리에서 로드)
CUSTOM_VOICE_PRESETS = {}

VOICES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "voices")


def _convert_to_wav(src_path, voice_dir):
    """비-wav 오디오 파일을 wav로 변환 (앞쪽 무음 제거 포함)"""
    import numpy as np
    try:
        import av
        import soundfile as sf
    except ImportError:
        print(f"[warn] PyAV 또는 soundfile 미설치 - {src_path} 변환 불가")
        return None

    wav_path = os.path.join(voice_dir, "reference.wav")
    try:
        container = av.open(src_path)
        stream = container.streams.audio[0]
        sr = stream.rate or stream.codec_context.sample_rate

        frames = []
        for frame in container.decode(audio=0):
            arr = frame.to_ndarray()
            if arr.ndim > 1:
                arr = arr.mean(axis=0)
            frames.append(arr)
        container.close()

        audio = np.concatenate(frames).astype(np.float32)
        if audio.max() > 1.0 or audio.min() < -1.0:
            audio = audio / 32768.0

        # 앞쪽 무음 제거
        threshold = 0.01
        abs_audio = np.abs(audio)
        window = int(sr * 0.02)
        smoothed = np.convolve(abs_audio, np.ones(window) / window, mode='same')
        nonsilent = np.where(smoothed > threshold)[0]
        if len(nonsilent) > 0:
            start = max(0, nonsilent[0] - int(sr * 0.05))
            audio = audio[start:]

        sf.write(wav_path, audio, sr)
        print(f"[info] 변환 완료: {os.path.basename(src_path)} -> reference.wav ({len(audio)/sr:.1f}s)")
        return wav_path
    except Exception as e:
        print(f"[warn] 오디오 변환 실패 ({src_path}): {e}")
        return None


def load_custom_voices():
    """voices/ 디렉토리에서 커스텀 음성 프리셋을 로드"""
    CUSTOM_VOICE_PRESETS.clear()

    if not os.path.isdir(VOICES_DIR):
        return CUSTOM_VOICE_PRESETS

    for voice_dir_name in os.listdir(VOICES_DIR):
        voice_dir = os.path.join(VOICES_DIR, voice_dir_name)
        preset_path = os.path.join(voice_dir, "preset.json")

        if not os.path.isdir(voice_dir) or not os.path.isfile(preset_path):
            continue

        try:
            with open(preset_path, 'r', encoding='utf-8') as f:
                preset = json.load(f)
        except Exception:
            continue

        # 참조 오디오 파일 찾기
        ref_audio_path = None
        # preset.json에 ref_audio 필드가 있으면 우선 사용
        if preset.get("ref_audio"):
            candidate = os.path.join(voice_dir, preset["ref_audio"])
            if os.path.isfile(candidate):
                ref_audio_path = candidate
        # 없으면 reference.* 패턴으로 폴백
        if ref_audio_path is None:
            for ext in ('.wav', '.m4a', '.mp3', '.flac', '.ogg'):
                candidate = os.path.join(voice_dir, f"reference{ext}")
                if os.path.isfile(candidate):
                    ref_audio_path = candidate
                    break

        if ref_audio_path is None:
            continue

        # 비-wav 파일이면 wav로 자동 변환
        if ref_audio_path and not ref_audio_path.lower().endswith('.wav'):
            converted = _convert_to_wav(ref_audio_path, voice_dir)
            if converted:
                ref_audio_path = converted
            else:
                continue  # 변환 실패 시 건너뜀

        voice_name = preset.get("name", voice_dir_name)
        CUSTOM_VOICE_PRESETS[voice_name] = {
            "ref_audio_path": ref_audio_path,
            "ref_text": preset.get("ref_text", ""),
            "lang_code": preset.get("lang_code", "ko"),
            "language": preset.get("language", "Korean"),
            "description": preset.get("description", "커스텀 보이스 클론"),
            "is_custom": True,
        }

    return CUSTOM_VOICE_PRESETS


def get_all_voice_options():
    """built-in + custom 음성 목록을 합쳐서 반환"""
    all_voices = list(VOICE_OPTIONS)
    for voice_name, info in CUSTOM_VOICE_PRESETS.items():
        all_voices.append((voice_name, info["language"], info["description"]))
    return all_voices
